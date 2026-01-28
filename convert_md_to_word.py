#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将包含Mermaid图表的Markdown文档转换为Word文档
使用Mermaid Live Editor在线API生成图表图片
"""

import os
import re
import base64
import requests
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from io import BytesIO


class MarkdownToWordConverter:
    def __init__(self, md_file_path, output_word_path=None):
        """
        初始化转换器

        Args:
            md_file_path: Markdown文件路径
            output_word_path: 输出的Word文件路径(可选)
        """
        self.md_file_path = md_file_path
        self.output_word_path = output_word_path or self._generate_output_path()
        self.temp_image_dir = "temp_images"
        self.mermaid_api_url = "https://mermaid.ink/svg/{encoded_code}"

        # 创建临时图片目录
        os.makedirs(self.temp_image_dir, exist_ok=True)

    def _generate_output_path(self):
        """生成输出文件路径"""
        base_name = os.path.splitext(os.path.basename(self.md_file_path))[0]
        dir_name = os.path.dirname(self.md_file_path)
        return os.path.join(dir_name, f"{base_name}.docx")

    def _read_markdown(self):
        """读取Markdown文件"""
        with open(self.md_file_path, 'r', encoding='utf-8') as f:
            return f.read()

    def _extract_mermaid_blocks(self, content):
        """
        提取Mermaid代码块

        Returns:
            list: [(start_pos, end_pos, mermaid_code), ...]
        """
        pattern = r'```mermaid\n(.*?)\n```'
        matches = list(re.finditer(pattern, content, re.DOTALL))
        return [(m.start(), m.end(), m.group(1).strip()) for m in matches]

    def _convert_mermaid_to_svg(self, mermaid_code, index):
        """
        使用在线API将Mermaid代码转换为SVG

        Args:
            mermaid_code: Mermaid代码
            index: 图表索引

        Returns:
            str: SVG文件路径,如果失败返回None
        """
        try:
            # 编码Mermaid代码
            encoded = base64.urlsafe_b64encode(mermaid_code.encode()).decode()
            url = self.mermaid_api_url.format(encoded_code=encoded)

            # 下载SVG
            response = requests.get(url, timeout=30)
            response.raise_for_status()

            svg_path = os.path.join(self.temp_image_dir, f"mermaid_{index}.svg")
            with open(svg_path, 'w', encoding='utf-8') as f:
                f.write(response.text)

            print(f"✓ 成功生成图表 {index + 1}: {os.path.basename(svg_path)}")
            return svg_path

        except Exception as e:
            print(f"✗ 生成图表 {index + 1} 失败: {str(e)}")
            return None

    def _parse_markdown_to_blocks(self, content, mermaid_blocks):
        """
        将Markdown解析为文本块和Mermaid块的交替序列

        Returns:
            list: [{"type": "text"/"mermaid", "content": "...", "index": int}, ...]
        """
        blocks = []
        last_end = 0

        for i, (start, end, mermaid_code) in enumerate(mermaid_blocks):
            # 添加Mermaid之前的文本
            if start > last_end:
                text = content[last_end:start]
                if text.strip():
                    blocks.append({"type": "text", "content": text})

            # 添加Mermaid块
            blocks.append({
                "type": "mermaid",
                "content": mermaid_code,
                "index": i
            })

            last_end = end

        # 添加最后的文本
        if last_end < len(content):
            text = content[last_end:]
            if text.strip():
                blocks.append({"type": "text", "content": text})

        return blocks

    def _add_heading(self, doc, text, level=1):
        """添加标题"""
        heading = doc.add_heading(text, level=level)
        # 设置中文字体
        for run in heading.runs:
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        return heading

    def _add_paragraph(self, doc, text, style=None):
        """添加段落"""
        para = doc.add_paragraph(text, style=style)
        # 设置中文字体
        for run in para.runs:
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        return para

    def _parse_heading_level(self, text):
        """解析标题级别"""
        match = re.match(r'^(#{1,6})\s+(.*)', text)
        if match:
            level = len(match.group(1))
            title = match.group(2).strip()
            return level, title
        return None, None

    def _process_text_block(self, doc, text):
        """
        处理文本块

        Args:
            doc: Word文档对象
            text: 文本内容
        """
        lines = text.split('\n')
        current_para_lines = []

        for line in lines:
            # 跳过空行
            if not line.strip():
                if current_para_lines:
                    self._add_paragraph(doc, '\n'.join(current_para_lines))
                    current_para_lines = []
                continue

            # 检查是否是标题
            level, title = self._parse_heading_level(line)
            if level:
                # 先添加之前的段落
                if current_para_lines:
                    self._add_paragraph(doc, '\n'.join(current_para_lines))
                    current_para_lines = []
                # 添加标题
                self._add_heading(doc, title, level=min(level, 3))  # Word最多支持3级
                continue

            # 检查是否是分隔线
            if line.strip() == '---':
                if current_para_lines:
                    self._add_paragraph(doc, '\n'.join(current_para_lines))
                    current_para_lines = []
                # 添加分隔线
                p = doc.add_paragraph()
                p.add_run('_' * 80).bold = True
                continue

            # 检查是否是列表
            if line.strip().startswith(('- ', '* ', '• ')) or re.match(r'^\d+\.\s', line):
                if current_para_lines:
                    self._add_paragraph(doc, '\n'.join(current_para_lines))
                    current_para_lines = []
                # 添加列表项
                self._add_paragraph(doc, line)
                continue

            # 普通文本行
            current_para_lines.append(line)

        # 添加最后的段落
        if current_para_lines:
            self._add_paragraph(doc, '\n'.join(current_para_lines))

    def _add_image(self, doc, svg_path, width=Inches(6)):
        """
        在文档中添加图片

        注意: python-docx不直接支持SVG,这里需要先将SVG转换为PNG
        为了简化,我们只添加图片占位符
        """
        try:
            # 在实际应用中,应该将SVG转换为PNG
            # 这里使用cairosvg或svg2png等工具
            import cairosvg

            png_path = svg_path.replace('.svg', '.png')
            cairosvg.svg2png(url=svg_path, write_to=png_path, output_width=1200)

            # 添加PNG图片
            doc.add_picture(png_path, width=width)
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            print(f"✓ 添加图片: {os.path.basename(png_path)}")
            return True

        except ImportError:
            print("⚠ 警告: 未安装cairosvg库,无法添加图片")
            print("  请运行: pip install cairosvg")
            # 添加文本占位符
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(f"[图表图片: {os.path.basename(svg_path)}]").bold = True
            p.add_run(f"\n(文件位置: {svg_path})")
            return False

        except Exception as e:
            print(f"✗ 添加图片失败: {str(e)}")
            return False

    def _setup_document_styles(self, doc):
        """设置文档样式"""
        # 设置默认字体
        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(11)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    def convert(self):
        """
        执行转换

        Returns:
            str: 生成的Word文件路径
        """
        print(f"开始转换: {self.md_file_path}")
        print("-" * 60)

        # 读取Markdown
        content = self._read_markdown()
        print(f"✓ 读取Markdown文件: {len(content)} 字符")

        # 提取Mermaid块
        mermaid_blocks = self._extract_mermaid_blocks(content)
        print(f"✓ 发现 {len(mermaid_blocks)} 个Mermaid图表")

        # 生成Mermaid SVG图片
        svg_files = {}
        for i, (start, end, mermaid_code) in enumerate(mermaid_blocks):
            svg_path = self._convert_mermaid_to_svg(mermaid_code, i)
            if svg_path:
                svg_files[i] = svg_path

        # 解析Markdown为块序列
        blocks = self._parse_markdown_to_blocks(content, mermaid_blocks)

        # 创建Word文档
        doc = Document()
        self._setup_document_styles(doc)

        # 添加文档标题
        title_match = re.search(r'^#\s+(.+)', content, re.MULTILINE)
        if title_match:
            self._add_heading(doc, title_match.group(1), level=0)

        # 处理每个块
        print("-" * 60)
        print("生成Word文档...")
        for block in blocks:
            if block['type'] == 'text':
                self._process_text_block(doc, block['content'])
            elif block['type'] == 'mermaid':
                index = block['index']
                if index in svg_files:
                    self._add_image(doc, svg_files[index])

        # 保存Word文档
        doc.save(self.output_word_path)
        print("-" * 60)
        print(f"✓ 转换完成! 输出文件: {self.output_word_path}")

        return self.output_word_path

    def cleanup(self):
        """清理临时文件"""
        if os.path.exists(self.temp_image_dir):
            import shutil
            shutil.rmtree(self.temp_image_dir)
            print(f"✓ 清理临时文件: {self.temp_image_dir}")


def main():
    """主函数"""
    import sys

    # 检查命令行参数
    if len(sys.argv) > 1:
        md_file = sys.argv[1]
    else:
        # 默认使用COMPLETE_ARCHITECTURE.md
        md_file = "docs/COMPLETE_ARCHITECTURE.md"

    # 检查文件是否存在
    if not os.path.exists(md_file):
        print(f"错误: 文件不存在 - {md_file}")
        print("用法: python convert_md_to_word.py <markdown文件路径>")
        sys.exit(1)

    try:
        # 创建转换器并执行转换
        converter = MarkdownToWordConverter(md_file)
        output_file = converter.convert()

        print("\n" + "=" * 60)
        print("转换成功!")
        print(f"输入文件: {md_file}")
        print(f"输出文件: {output_file}")
        print("=" * 60)

        # 询问是否清理临时文件
        response = input("\n是否删除临时SVG文件? (y/n): ").strip().lower()
        if response == 'y':
            converter.cleanup()

    except KeyboardInterrupt:
        print("\n\n转换已取消")
    except Exception as e:
        print(f"\n错误: {str(e)}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()
