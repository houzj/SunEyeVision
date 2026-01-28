import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
import requests

"""
Renders mermaid (.mmd) code blocks found in a Markdown file using Kroki (https://kroki.io/)
and embeds the rendered images into a Word (.docx) document.

Usage:
    python render_and_embed_mermaid.py <input.md> [output.docx]

Notes:
- Requires internet access to call Kroki.
- Installs required packages: requests, cairosvg
"""


def render_mermaid_to_png(mermaid_source: str, out_png: Path) -> None:
    # Use Kroki to render mermaid directly to PNG to avoid native cairo dependency
    kroki_url = 'https://kroki.io/mermaid/png'
    # request PNG with a timeout
    r = requests.post(kroki_url, data=mermaid_source.encode('utf-8'), timeout=30)
    if r.status_code != 200:
        raise RuntimeError(f'Kroki render failed: {r.status_code} {r.text[:200]}')
    with out_png.open('wb') as outf:
        outf.write(r.content)


def convert_and_embed(md_path: Path, out_docx: Path, image_width_inch: float = 6.0):
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    in_code = False
    code_lines = []
    current_lang = ''
    mermaid_count = 0

    out_dir = md_path.parent
    mermaid_dir = out_dir / 'mermaid_blocks'
    mermaid_dir.mkdir(parents=True, exist_ok=True)

    with md_path.open('r', encoding='utf-8') as f:
        for raw in f:
            line = raw.rstrip('\n')
            stripped_line = line.strip()

            if stripped_line.startswith('```'):
                lang = stripped_line[3:].strip()
                if not in_code:
                    in_code = True
                    current_lang = lang.lower()
                    code_lines = []
                else:
                    in_code = False
                    if current_lang == 'mermaid':
                        mermaid_count += 1
                        name = f'mermaid_{mermaid_count}.mmd'
                        mmd_path = mermaid_dir / name
                        with mmd_path.open('w', encoding='utf-8') as outf:
                            outf.write('\n'.join(code_lines) + '\n')

                        png_name = f'mermaid_{mermaid_count}.png'
                        png_path = mermaid_dir / png_name
                        print(f'Rendering {mmd_path} -> {png_path} via Kroki...')
                        # skip rendering if PNG already exists
                        if not png_path.exists():
                            try:
                                render_mermaid_to_png('\n'.join(code_lines), png_path)
                            except Exception as e:
                                print('Rendering failed:', e)
                                # fallback: include source as code block
                                p = doc.add_paragraph()
                                p.add_run(f'[Failed to render mermaid (see {mmd_path.as_posix()}): {e}]')
                                # include source
                                p2 = doc.add_paragraph()
                                p2.add_run('\n'.join(code_lines))
                                code_lines = []
                                current_lang = ''
                                continue
                        else:
                            print(f'PNG already exists, skipping rendering: {png_path}')

                        # insert image into doc
                        try:
                            doc.add_picture(str(png_path), width=Inches(image_width_inch))
                        except Exception as e:
                            print('Failed to insert image into DOCX:', e)
                            p = doc.add_paragraph()
                            p.add_run(f'[Rendered image saved to {png_path.as_posix()} but failed to embed: {e}]')

                        # optional: add a small caption with filename
                        cap = doc.add_paragraph()
                        cap_run = cap.add_run(png_name)
                        cap_run.italic = True

                    else:
                        # generic code block
                        p = doc.add_paragraph()
                        run = p.add_run('\n'.join(code_lines))
                        run.font.name = 'Consolas'
                        run.font.size = Pt(9)

                    code_lines = []
                    current_lang = ''
                continue

            if in_code:
                code_lines.append(line)
                continue

            # headings
            stripped = line.lstrip()
            if stripped.startswith('#'):
                hashes = len(stripped) - len(stripped.lstrip('#'))
                text = stripped.lstrip('#').strip()
                if text == '':
                    continue
                level = min(hashes, 4)
                doc.add_heading(text, level=level)
                continue

            # lists (simple)
            s = line.lstrip()
            if s.startswith('- ') or s.startswith('* ') or s.startswith('+ '):
                text = s[2:].strip()
                doc.add_paragraph(text, style='List Bullet')
                continue

            # blank line
            if line.strip() == '':
                doc.add_paragraph('')
                continue

            # normal paragraph
            doc.add_paragraph(line)

    # handle unclosed code block
    if in_code and code_lines:
        if current_lang == 'mermaid':
            mermaid_count += 1
            name = f'mermaid_{mermaid_count}.mmd'
            mmd_path = mermaid_dir / name
            with mmd_path.open('w', encoding='utf-8') as outf:
                outf.write('\n'.join(code_lines) + '\n')
            png_path = mermaid_dir / f'mermaid_{mermaid_count}.png'
            try:
                render_mermaid_to_png('\n'.join(code_lines), png_path)
                doc.add_picture(str(png_path), width=Inches(image_width_inch))
            except Exception as e:
                p = doc.add_paragraph()
                p.add_run(f'[Failed to render mermaid: {e}]')
        else:
            p = doc.add_paragraph()
            p.add_run('\n'.join(code_lines))

    doc.save(out_docx)
    print('Saved', out_docx)


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('Usage: python render_and_embed_mermaid.py <input.md> [output.docx]')
        sys.exit(1)
    md = Path(sys.argv[1])
    if not md.exists():
        print('Input markdown not found:', md)
        sys.exit(2)
    out = Path(sys.argv[2]) if len(sys.argv) > 2 else md.with_suffix('.docx')
    convert_and_embed(md, out)
