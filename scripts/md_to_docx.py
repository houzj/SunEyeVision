import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

# Markdown -> DOCX converter with special handling for mermaid blocks.
# Mermaid blocks are exported to docs/mermaid_blocks/*.mmd and the DOCX
# will contain a placeholder paragraph referencing each generated .mmd file.
# Usage: python md_to_docx.py <input.md> [output.docx]


def add_code_paragraph(doc, code_text, font_name='Consolas', size=9):
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    # set monospace font if possible
    try:
        run.font.name = font_name
        r = run._element.rPr
        r.rFonts.set(qn('w:eastAsia'), font_name)
    except Exception:
        pass
    run.font.size = Pt(size)


def convert(md_path, out_path):
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    in_code = False
    code_lines = []
    current_lang = ''
    mermaid_count = 0

    out_dir = md_path.parent
    mermaid_dir = out_dir / 'mermaid_blocks'
    mermaid_dir.mkdir(parents=True, exist_ok=True)

    with md_path.open('r', encoding='utf-8') as f:
        for raw in f:
            line = raw.rstrip('\n')
            stripped_line = line.strip()

            if stripped_line.startswith('```'):
                # parse language (if any)
                lang = stripped_line[3:].strip()
                if not in_code:
                    # enter code block
                    in_code = True
                    current_lang = lang.lower()
                    code_lines = []
                else:
                    # close code block
                    in_code = False
                    if current_lang == 'mermaid':
                        mermaid_count += 1
                        name = f"mermaid_{mermaid_count}.mmd"
                        target = mermaid_dir / name
                        # write the mermaid source
                        with target.open('w', encoding='utf-8') as outf:
                            outf.write('\n'.join(code_lines) + '\n')

                        # add placeholder paragraph to docx
                        p = doc.add_paragraph()
                        p.add_run(f"[Mermaid diagram saved to: {target.as_posix()}]")
                        # also include the mermaid source as a small code block for reference
                        add_code_paragraph(doc, '\n'.join(code_lines), size=9)
                    else:
                        # generic code block: insert as monospace paragraph
                        add_code_paragraph(doc, '\n'.join(code_lines), size=9)

                    code_lines = []
                    current_lang = ''
                continue

            if in_code:
                code_lines.append(line)
                continue

            # headings
            stripped = line.lstrip()
            if stripped.startswith('#'):
                hashes = len(stripped) - len(stripped.lstrip('#'))
                text = stripped.lstrip('#').strip()
                if text == '':
                    continue
                level = min(hashes, 4)
                doc.add_heading(text, level=level)
                continue

            # lists (simple, only top-level bullets)
            s = line.lstrip()
            if s.startswith('- ') or s.startswith('* ') or s.startswith('+ '):
                text = s[2:].strip()
                p = doc.add_paragraph(text, style='List Bullet')
                continue

            if s.startswith('1. ') or (len(s) > 2 and s[0:2].isdigit() and s[2:].strip().startswith('. ')):
                # crude numeric list handling -> fallback to normal paragraph
                doc.add_paragraph(line)
                continue

            # blank line -> add an empty paragraph to preserve spacing
            if line.strip() == '':
                doc.add_paragraph('')
                continue

            # normal paragraph
            doc.add_paragraph(line)

    # if file ended while still in code block
    if in_code and code_lines:
        if current_lang == 'mermaid':
            mermaid_count += 1
            name = f"mermaid_{mermaid_count}.mmd"
            target = mermaid_dir / name
            with target.open('w', encoding='utf-8') as out_f:
                out_f.write('\n'.join(code_lines) + '\n')
            p = doc.add_paragraph()
            p.add_run(f"[Mermaid diagram saved to: {target.as_posix()}]")
            add_code_paragraph(doc, '\n'.join(code_lines), size=9)
        else:
            add_code_paragraph(doc, '\n'.join(code_lines), size=9)

    doc.save(out_path)


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('Usage: python md_to_docx.py <input.md> [output.docx]')
        sys.exit(1)
    md = Path(sys.argv[1])
    if not md.exists():
        print('Input markdown not found:', md)
        sys.exit(2)
    out = Path(sys.argv[2]) if len(sys.argv) > 2 else md.with_suffix('.docx')
    convert(md, out)
    print('Saved', out)
